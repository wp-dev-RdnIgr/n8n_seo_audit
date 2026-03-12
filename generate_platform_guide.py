#!/usr/bin/env python3
"""Генерація лаконічної інструкції платформи SEO Audit (укр) у DOCX."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# -- Поля сторінки --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Стилі --
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.paragraph_format.space_after = Pt(4)
style_normal.paragraph_format.line_spacing = 1.15

for lvl in range(1, 4):
    hs = doc.styles[f'Heading {lvl}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)


# -- Хелпери --
def tip(text):
    p = doc.add_paragraph()
    r = p.add_run(f"💡 {text}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)
    r.font.italic = True

def warn(text):
    p = doc.add_paragraph()
    r = p.add_run(f"⚠️ {text}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    r.font.italic = True

def step(n, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{n}. ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    p.add_run(text)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(f" — {text}")
    else:
        p.add_run(text)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light List Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row_data in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
    return t

def section_break():
    doc.add_paragraph()


# =============================================
# ТИТУЛЬНА СТОРІНКА
# =============================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run("SEO Audit Platform")
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
r.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Інструкція користувача")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(16)
r3 = p3.add_run(
    "SimilarWeb Comparing  •  SEO Audit Wizard  •  Email Reports"
)
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(30)
r4 = p4.add_run("v2.0 | Березень 2026")
r4.font.size = Pt(10)
r4.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# =============================================
# ЗМІСТ
# =============================================
doc.add_heading("Зміст", level=1)
toc = [
    "Частина 1. SimilarWeb Comparing",
    "  1.1. Клієнти та конкуренти",
    "  1.2. Збір даних",
    "  1.3. Dashboard",
    "  1.4. Аналіз (Reviews)",
    "  1.5. Черга та логи",
    "  1.6. Налаштування",
    "Частина 2. SEO Audit Wizard",
    "  2.1. Базова інформація та конкуренти",
    "  2.2. Семантика та ключові слова",
    "  2.3. PageSpeed",
    "  2.4. AI-аналіз конкурентів",
    "  2.5. PDF-парсинг",
    "  2.6. Технічний аудит (Screaming Frog)",
    "Частина 3. Email-звіти",
    "FAQ",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(1)
    if item.startswith("  "):
        p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# =============================================
# ЧАСТИНА 1: SIMILARWEB COMPARING
# =============================================
doc.add_heading("Частина 1. SimilarWeb Comparing", level=1)
doc.add_paragraph(
    "Панель порівняння трафіку клієнтського сайту з конкурентами. "
    "Дані збираються автоматично з SimilarWeb, зберігаються в Supabase."
)

# -- 1.1 Клієнти --
doc.add_heading("1.1. Клієнти та конкуренти", level=2)

doc.add_heading("Додати клієнта:", level=3)
step(1, 'Вкладка "Clients" → поле "Client Domain" → введіть домен')
step(2, '"Manager Contact" — email або @telegram (необов\'язково)')
step(3, 'Натисніть "+ Add"')
tip("Домен нормалізується автоматично: видаляється https://, www., слеш.")
section_break()

doc.add_heading("Додати конкурентів:", level=3)
step(1, 'У рядку клієнта натисніть 👥')
step(2, "У модальному вікні введіть домен конкурента → «+ Add»")
step(3, "Повторіть для кожного конкурента")
tip("Оптимально: 5–15 конкурентів на клієнта.")
section_break()

doc.add_heading("Видалення:", level=3)
doc.add_paragraph(
    "Клієнт — іконка 🗑 у рядку. Конкурент — ✕ у модальному вікні. "
    "Видалення м'яке (soft delete): дані зберігаються в БД."
)

doc.add_page_break()

# -- 1.2 Збір даних --
doc.add_heading("1.2. Збір даних", level=2)

doc.add_heading("Автоматичний спосіб (рекомендовано):", level=3)
doc.add_paragraph(
    'Dashboard → кнопка "Build Queue". Система сама визначить відсутні періоди '
    "та створить задачі. Збір: 5–30 хв залежно від кількості сайтів."
)

doc.add_heading("Ручний спосіб:", level=3)
doc.add_paragraph(
    'Queue → обрати клієнта → обрати місяць → "Create". '
    "Дублікати пропускаються автоматично."
)

add_table(
    ["Статус", "Значення"],
    [
        ("Pending", "Очікує обробки"),
        ("Processing", "Робот збирає дані"),
        ("Done", "Дані зібрані"),
        ("Error", "Помилка — перевірте Logs"),
    ]
)
section_break()

doc.add_heading("Статуси клієнтів:", level=3)
add_table(
    ["Іконка", "Статус", "Що означає"],
    [
        ("🔄", "Updating", "Є активні задачі, дані збираються"),
        ("✅", "Done", "Всі задачі завершені"),
        ("⏳", "New", "Задачі ще не створювались"),
    ]
)

doc.add_page_break()

# -- 1.3 Dashboard --
doc.add_heading("1.3. Dashboard", level=2)
doc.add_paragraph("Стартова сторінка зі зведеною статистикою.")

doc.add_heading("Картки:", level=3)
add_table(
    ["Картка", "Що показує"],
    [
        ("Clients", "Кількість активних клієнтів"),
        ("Competitors", "Загальна кількість конкурентів"),
        ("Pending / Done", "Задачі в черзі / завершені"),
        ("Errors", "Нерозв'язані помилки"),
        ("Data Rows", "Записів даних у БД"),
    ]
)
section_break()

doc.add_heading("Кнопки:", level=3)
bullet("Додати клієнта", "Add Client")
bullet("Перейти до черги", "Queue")
bullet("Оновити дані на сторінці", "Refresh")
bullet("Автоматично створити задачі для відсутніх даних", "Build Queue")
section_break()

doc.add_paragraph(
    'Таблиця "Client Reviews" — список клієнтів зі статусом та кнопкою "Review" '
    "для переходу до аналізу."
)

doc.add_page_break()

# -- 1.4 Reviews --
doc.add_heading("1.4. Аналіз (Reviews)", level=2)
doc.add_paragraph(
    "Основний інструмент аналізу. 3 під-вкладки:"
)

doc.add_heading("Client Overview", level=3)
bullet("Порівняння з попереднім місяцем (% зміни, кольорова індикація)", "MoM")
bullet("Порівняння з тим же місяцем минулого року", "YoY")
bullet("Повна таблиця за обраний діапазон дат", "All Metrics")
section_break()

doc.add_heading("Competitor Comparison", level=3)
bullet("Стовпчиковий графік: клієнт vs конкуренти за обрану метрику", "MoM / YoY Chart")
bullet("Всі метрики × всі сайти. Клієнт виділений синім. 🏆 = найкращий показник", "Summary Table")
tip("Мультиселект дозволяє обрати, яких конкурентів показувати.")
section_break()

doc.add_heading("Traffic Sources", level=3)
bullet("Горизонтальні бари трафіку + таблиця залученості", "Traffic & Engagement")
bullet("Лінійний графік обраної метрики за весь період", "Traffic Over Time")
bullet("7 вкладок: Direct, Organic, Paid, Display, Social, Email, AI", "Channel Traffic")
bullet("Кругова діаграма розподілу каналів", "Source Distribution")
bullet("Графік зміни пропорцій каналів з часом", "Source Trends")
bullet("Числові значення по кожному каналу. 🏆 = лідер", "Source Details")

doc.add_page_break()

# -- 1.5 Метрики --
doc.add_heading("Доступні метрики (15):", level=2)
add_table(
    ["Метрика", "Опис"],
    [
        ("Monthly Visits", "Загальна кількість візитів"),
        ("Unique Visitors", "Унікальні відвідувачі"),
        ("Visits per Visitor", "Візитів на відвідувача"),
        ("Deduplicated Audience", "Дедуплікована аудиторія"),
        ("Page Views", "Перегляди сторінок"),
        ("Visit Duration", "Тривалість візиту"),
        ("Pages per Visit", "Сторінок за візит"),
        ("Bounce Rate", "Показник відмов (%)"),
        ("Direct", "Прямий трафік"),
        ("Organic Search", "Органічний пошук"),
        ("Paid Search", "Платний пошук"),
        ("Display Ads", "Медійна реклама"),
        ("Social", "Соціальні мережі"),
        ("Email", "Email-розсилки"),
        ("AI Traffic", "AI-сервіси"),
    ]
)

doc.add_page_break()

# -- 1.6 Queue & Logs --
doc.add_heading("1.5. Черга та логи", level=2)

doc.add_heading("Queue (черга задач):", level=3)
doc.add_paragraph(
    "Фільтр за статусом: All / Pending / Processing / Done / Error. "
    "Дії: Retry (перезапустити), Delete (видалити)."
)
warn("Не видаляйте задачі зі статусом Processing.")
section_break()

doc.add_heading("Logs (журнал помилок):", level=3)
doc.add_paragraph(
    "Фільтр: All / Unresolved / Resolved. "
    "Кожен запис: дата, тип робота, період, сайти, текст помилки, кількість спроб."
)
bullet("Перезапуск відповідної задачі", "Retry")
bullet("Позначити як розв'язану", "Mark Resolved")

doc.add_page_break()

# -- 1.7 Settings --
doc.add_heading("1.6. Налаштування", level=2)
doc.add_paragraph(
    "Supabase URL та Service Role Key. Кнопки: Save, Test Connection."
)
warn("Не змінюйте без узгодження з адміністратором!")

doc.add_page_break()

# =============================================
# ЧАСТИНА 2: SEO AUDIT WIZARD
# =============================================
doc.add_heading("Частина 2. SEO Audit Wizard", level=1)
doc.add_paragraph(
    "Покроковий майстер для комплексного SEO-аудиту. "
    "Бічна панель з 9 кроками, розділеними на 3 етапи."
)

add_table(
    ["Етап", "Кроки"],
    [
        ("1. Збір даних", "Basic Info → Competitors → Semantic → Keywords → PageSpeed → Run"),
        ("2. Обробка", "AI Analysis → PDF Parsing"),
        ("3. Тех. аудит", "Technical SEO Audit"),
    ]
)
section_break()

# -- 2.1 --
doc.add_heading("2.1. Базова інформація та конкуренти", level=2)

doc.add_heading("Basic Info (крок 1):", level=3)
step(1, "Введіть email менеджера (обов'язково)")
step(2, "Введіть домен клієнта")
tip("Домен нормалізується автоматично, як і в SimilarWeb Comparing.")
section_break()

doc.add_heading("Competitors (крок 2):", level=3)
doc.add_paragraph(
    "Введіть домен конкурента → Enter для додавання. "
    "Backspace — видалити останній. Ctrl+V — вставити кілька одразу. "
    "Кожен конкурент відображається як тег."
)

doc.add_page_break()

# -- 2.2 --
doc.add_heading("2.2. Семантика та ключові слова", level=2)

doc.add_heading("Semantic Expansion (крок 3):", level=3)
doc.add_paragraph(
    "URL Google Sheets з seed-фразами (перша колонка, перший аркуш). "
    "Необов'язковий крок. Можна вказати папку Google Drive для результатів."
)

doc.add_heading("Keyword Metrics (крок 4):", level=3)
doc.add_paragraph(
    "URL Google Sheets зі списком ключових слів. "
    "Необов'язковий крок. Формат аналогічний."
)

# -- 2.3 --
doc.add_heading("2.3. PageSpeed (крок 5)", level=2)
doc.add_paragraph(
    "URL Google Sheets для аналізу PageSpeed. "
    "Можна обрати метрики та вказати папку для результатів. Необов'язковий крок."
)

# -- 2.4 --
doc.add_heading("2.4. AI-аналіз конкурентів (крок 7)", level=2)
doc.add_paragraph(
    "Для кожного конкурента вкажіть URL Google Sheets зі зібраними даними. "
    "AI аналізує всі таблиці та створює ОДИН Google Doc зі зведеним звітом."
)
tip("Час обробки: ~2–4 хвилини на конкурента.")
warn("Не закривайте сторінку під час обробки!")

doc.add_page_break()

# -- 2.5 --
doc.add_heading("2.5. PDF-парсинг (крок 8)", level=2)
doc.add_paragraph(
    "Завантажте URL PDF-файлу з Google Drive. "
    "Система розпарсить PDF та створить структурований Google Sheets. "
    "Результат: URL таблиці, кількість аркушів, час обробки."
)

# -- 2.6 --
doc.add_heading("2.6. Технічний аудит — Screaming Frog (крок 9)", level=2)

doc.add_paragraph("Аналіз технічного стану сайту на базі експорту Screaming Frog.")

doc.add_heading("Вхідні дані:", level=3)
bullet("Домен сайту (обов'язково)")
bullet("internal_all.xlsx — експорт зі Screaming Frog (обов'язково)")
bullet("Файли images alt — зображення без ALT-тексту (необов'язково)")
bullet("Файл бєклінків (необов'язково)")
bullet("Email менеджера та URL папки (необов'язково)")

doc.add_heading("Результат:", level=3)
bullet("Excel-звіт з технічним аудитом")
bullet("Google Doc зі зведеним аналізом")
bullet("Папка з усіма файлами")

doc.add_page_break()

# =============================================
# ЧАСТИНА 3: EMAIL ЗВІТИ
# =============================================
doc.add_heading("Частина 3. Email-звіти", level=1)
doc.add_paragraph(
    "Щомісячна автоматична розсилка зведених звітів SimilarWeb менеджерам. "
    "Звіти містять ключові метрики та тренди за місяць."
)
doc.add_paragraph(
    "Розсилка запускається автоматично або вручну через n8n."
)

doc.add_page_break()

# =============================================
# FAQ
# =============================================
doc.add_heading("FAQ", level=1)

faqs = [
    ("Як довго збираються дані?",
     "5–30 хвилин залежно від кількості сайтів. Статус — на Dashboard."),
    ("Немає даних у конкурента?",
     "SimilarWeb може не мати даних для малих сайтів. Перевірте Logs."),
    ("Що означає статус «New»?",
     "Задачі ще не створювались. Натисніть Build Queue."),
    ("Можу відновити видаленого конкурента?",
     "Так — просто додайте його знову, система відновить автоматично."),
    ("Дані не відображаються в Reviews?",
     "Перевірте: 1) правильний місяць, 2) є зібрані дані, 3) Supabase підключено (зелена точка)."),
    ("Що означає 🏆?",
     "Найкращий показник серед усіх сайтів по цій метриці."),
    ("Скільки конкурентів додавати?",
     "Рекомендація: 5–15. Більше — довше збір та складніші графіки."),
    ("Як часто оновлювати дані?",
     "Раз на місяць. SimilarWeb оновлює дані щомісячно."),
    ("AI-аналіз завис?",
     "Зачекайте до 15 хвилин. Якщо не завершився — перезапустіть."),
    ("Де знайти результати технічного аудиту?",
     "Посилання на Excel, Doc та папку з'являються після завершення кроку 9."),
]

for q, a in faqs:
    p_q = doc.add_paragraph()
    r_q = p_q.add_run(f"❓ {q}")
    r_q.bold = True
    p_a = doc.add_paragraph(a)
    p_a.paragraph_format.left_indent = Cm(0.5)
    p_a.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# =============================================
# ЩО МОЖНА / НЕ МОЖНА
# =============================================
doc.add_heading("Правила роботи", level=1)

doc.add_heading("✅ Можна:", level=3)
for t in [
    "Додавати клієнтів, конкурентів у будь-який час",
    "Запускати Build Queue щомісяця",
    "Retry задачі з помилками",
    "Порівнювати дані за будь-який доступний період",
    "Пропускати необов'язкові кроки в SEO Audit Wizard",
]:
    bullet(t)

section_break()

doc.add_heading("❌ Не можна:", level=3)
for t in [
    "Змінювати Settings без адміністратора",
    "Видаляти задачі зі статусом Processing",
    "Закривати сторінку під час AI-аналізу або PDF-парсингу",
    "Очікувати миттєвий результат — збір даних потребує часу",
    "Додавати невалідні домени (без крапки, з пробілами)",
]:
    bullet(t)

# =============================================
# ФУТЕР
# =============================================
doc.add_paragraph()
p_f = doc.add_paragraph()
p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_f.add_run("— WebPromo | SEO Audit Platform v2.0 —")
r_f.font.size = Pt(9)
r_f.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# Зберегти
out = "/home/user/n8n_seo_audit/SEO_Audit_Platform_Guide_UA.docx"
doc.save(out)
print(f"✅ Документ збережено: {out}")
