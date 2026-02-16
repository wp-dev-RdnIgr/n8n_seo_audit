#!/usr/bin/env python3
"""Generate project documentation as a formatted .docx file in Ukrainian."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x57)

doc.styles['Heading 1'].font.size = Pt(20)
doc.styles['Heading 2'].font.size = Pt(15)
doc.styles['Heading 3'].font.size = Pt(12)

# ── Helpers ──
NAVY = RGBColor(0x1B, 0x3A, 0x57)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
LINK_BLUE = RGBColor(0x1A, 0x73, 0xE8)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def add_paragraph(text, bold=False, italic=False, color=None, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A57"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            # Alternate row shading
            if ri % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p._p.get_or_add_pPr().append(shading)
    return p

# ═══════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SEO AUDIT AUTOMATION PLATFORM')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = NAVY

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Технічна документація проекту')
run.font.size = Pt(16)
run.font.color.rgb = GRAY

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('WebPromo | SEO Department')
run.font.size = Pt(12)
run.font.color.rgb = GRAY

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Лютий 2026')
run.font.size = Pt(12)
run.font.color.rgb = GRAY

doc.add_page_break()

# ═══════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════

doc.add_heading('Зміст', level=1)
toc_items = [
    '1. Огляд проекту',
    '2. Сховища даних та коду',
    '3. Архітектура системи',
    '4. Веб-інтерфейс (GAS) — сторінки та маршрутизація',
    '5. Бекенд-функції (Код.gs)',
    '6. n8n Воркфлоу — детальний опис',
    '   6.1. SEO-аудит домену (основний)',
    '   6.2. Мастер-оркестратор "Аналіз домену"',
    '   6.3. AI-генератор звіту',
    '   6.4. PDF Audit Parser',
    '   6.5. GKP Universal System (семантика)',
    '   6.6. PageSpeed Test',
    '   6.7. Аудит посилального профілю',
    '7. Зведена таблиця ендпоінтів',
    '8. Зовнішні API та авторизація',
    '9. Структура Google Drive',
    '10. Структура файлів у репозиторії',
    '11. Нотатки щодо деплою',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith('   '):
        p.paragraph_format.left_indent = Cm(1.5)

doc.add_page_break()

# ═══════════════════════════════════════
# 1. OVERVIEW
# ═══════════════════════════════════════

doc.add_heading('1. Огляд проекту', level=1)

add_paragraph(
    'Автоматизована платформа для комплексного SEO-аналізу. '
    'Побудована на n8n (автоматизація воркфлоу), Google Apps Script (фронтенд + бекенд-проксі) '
    'та зовнішніх API (Ahrefs, SimilarWeb, Serpstat, Google Ads, OpenAI).'
)

add_paragraph('Основні можливості:', bold=True)
add_bullet('Комплексний SEO-аудит домену (трафік, ключові слова, беклінки, поведінкові метрики)')
add_bullet('AI-аналіз зібраних даних з генерацією звіту у Google Doc')
add_bullet('Парсинг PDF-аудитів у структуровані Google-таблиці')
add_bullet('Збір семантичного ядра через Google Keyword Planner (генерація ідей + метрики)')
add_bullet('Масове тестування швидкості сторінок через PageSpeed Insights')
add_bullet('Аудит посилального профілю (беклінки, анкори, DR)')
add_bullet('Мастер-оркестратор для запуску всіх аналізів з однієї форми')

# ═══════════════════════════════════════
# 2. STORAGE
# ═══════════════════════════════════════

doc.add_heading('2. Сховища даних та коду', level=1)

add_table(
    ['Ресурс', 'URL', 'Опис'],
    [
        ['Google Apps Script', 'https://script.google.com/home/projects/1UM5kbeqZwQUuv4fe-cz61xjaLxYSNMH7DkFW7ez8YmIn3MCNlQNBmcCL/edit', 'Фронтенд (HTML-форми) + бекенд (webhook-проксі). Деплоїться як Web App'],
        ['GitHub', 'https://github.com/wp-dev-RdnIgr/n8n_seo_audit', 'Git-репозиторій: GAS-код + JSON-експорти n8n воркфлоу'],
        ['Google Drive', 'https://drive.google.com/drive/u/0/folders/1A3Ak929G1c4XmZpPtI2FP4glrFE2-Bx2', 'Коренева папка для всіх згенерованих звітів'],
        ['Веб-інтерфейс', 'https://sites.google.com/web-promo.com.ua/seo-audit/', 'Google Sites сторінка з вбудованими GAS-формами через iframe'],
        ['Browse AI', 'https://dashboard.browse.ai/workspaces/web-promo-com-ua/robots', 'Боти-парсери для скрапінгу SimilarWeb'],
        ['n8n', 'https://n8n.rnd.webpromo.tools', 'Self-hosted n8n інстанс з усіма воркфлоу'],
    ],
    col_widths=[3.5, 7, 6]
)

# ═══════════════════════════════════════
# 3. ARCHITECTURE
# ═══════════════════════════════════════

doc.add_heading('3. Архітектура системи', level=1)

add_paragraph('Потік даних:', bold=True)

arch_lines = [
    '[Google Sites сторінка]',
    '    |',
    '    +-- iframe --> [GAS Web App (?page=...)]',
    '                       |',
    '                       +-- google.script.run --> [Код.gs бекенд-функції]',
    '                                                     |',
    '                                                     +-- UrlFetchApp.fetch() --> [n8n вебхуки]',
    '                                                                                      |',
    '                                                                                      +-- Ahrefs API (MCP)',
    '                                                                                      +-- Google Ads API',
    '                                                                                      +-- Serpstat API',
    '                                                                                      +-- Browse AI (SimilarWeb)',
    '                                                                                      +-- OpenAI GPT-4o',
    '                                                                                      +-- Google Drive/Sheets/Docs API',
    '                                                                                      +-- PageSpeed Insights API',
]
for line in arch_lines:
    add_code_block(line)

add_paragraph('')
add_paragraph('Принцип роботи:', bold=True)
add_bullet('Користувач відкриває Google Sites сторінку')
add_bullet('Сторінка містить iframe з GAS Web App (HTML-форма)')
add_bullet('Форма викликає бекенд-функцію через google.script.run')
add_bullet('Бекенд-функція відправляє POST-запит на n8n webhook')
add_bullet('n8n воркфлоу виконує всю роботу (збір даних, аналіз, генерація звітів)')
add_bullet('Результат повертається у відповіді вебхуку і відображається у формі')

# ═══════════════════════════════════════
# 4. FRONTEND PAGES
# ═══════════════════════════════════════

doc.add_heading('4. Веб-інтерфейс (GAS) — сторінки та маршрутизація', level=1)

add_paragraph(
    'Маршрутизація виконується через URL-параметр ?page=... у функції doGet(e) (Код.gs). '
    'За замовчуванням відкривається сторінка "audit".'
)

add_table(
    ['Параметр page', 'HTML-файл', 'Назва', 'Опис'],
    [
        ['audit (за замовч.)', 'form.html', 'Аналіз домену', 'Ввід домену -> повний SEO-аудит + кнопка AI-звіту'],
        ['master', 'analiz_domenu_form.html', 'Аналіз домену — Мастер', 'Оркестратор: 5 блоків (клієнт, конкуренти, семантика, метрики, PageSpeed)'],
        ['gkp', 'gkp_form.html', 'Семантичне ядро (GKP)', 'Legacy-форма: ручний ввід seed-ключів (теги)'],
        ['gkp_ideas', 'gkp_ideas.html', 'GKP: Генерація ідей', 'Етап 1 — генерація ідей з таблиці seed-фраз'],
        ['gkp_metrics', 'gkp_metrics.html', 'GKP: Метрики', 'Етап 2 — отримання обсягів пошуку з таблиці ключів'],
        ['pagespeed', 'pagespeed_form.html', 'PageSpeed Test', 'Масове тестування швидкості з таблиці URL-ів'],
        ['pdf_audit', 'pdf_audit_form.html', 'SEO AI Інструменти', 'Дві секції: AI-аналіз домену + парсинг PDF-аудиту'],
    ],
    col_widths=[3, 4, 3.5, 6]
)

# ═══════════════════════════════════════
# 5. BACKEND FUNCTIONS
# ═══════════════════════════════════════

doc.add_heading('5. Бекенд-функції (Код.gs)', level=1)

add_paragraph(
    'Усі бекенд-функції знаходяться у файлі Код.gs. Кожна функція приймає дані з HTML-форми, '
    'валідує їх і відправляє POST-запит на відповідний n8n webhook.'
)

# 5.1
doc.add_heading('5.1. submitAudit(domain)', level=3)
add_bullet('Викликається з: form.html (сторінка "Аналіз домену")')
add_bullet('Webhook: POST /webhook/seo-organic-traffic-v74')
add_bullet('Payload: { domain, country: "ua", top_pages_limit: 50 }')
add_bullet('Відповідь: { success, spreadsheetUrl, message }')
add_bullet('Дія: Запускає повний SEO-аудит домену. Повертає посилання на Google Sheet.')

# 5.2
doc.add_heading('5.2. submitAIAnalysis(spreadsheetUrl)', level=3)
add_bullet('Викликається з: form.html (секція AI) та pdf_audit_form.html (картка "AI аналіз домену")')
add_bullet('Webhook: POST /webhook-test/seo-audit-ai-report')
add_bullet('Payload: { url: spreadsheetUrl }')
add_bullet('Відповідь: { success, docUrl, domain, message }')
add_bullet('Дія: Надсилає таблицю на AI-аналіз. GPT-4o аналізує всі листи, генерує Google Doc.')

# 5.3
doc.add_heading('5.3. submitAnalizDomenu(formData)', level=3)
add_bullet('Викликається з: analiz_domenu_form.html (сторінка "Мастер")')
add_bullet('Webhook: POST /webhook/analiz-domenu')
add_bullet('Payload: { manager_email, client_domain?, competitors[]?, semantic_expansion?, metrics_collection?, pagespeed? }')
add_bullet('Відповідь: { success, folderUrl, details, message }')
add_bullet('Дія: Мастер-оркестратор. Створює структуру папок на Google Drive, послідовно запускає обрані блоки аналізу.')

# 5.4
doc.add_heading('5.4. submitGKP(formData) — legacy', level=3)
add_bullet('Викликається з: gkp_form.html')
add_bullet('Webhook: POST /webhook/gkp-ideas')
add_bullet('Payload: { doc_name, language, geo_target, seed_keywords, url_mapping, limit }')
add_bullet('Відповідь: { success, spreadsheetUrl, totalKeywords, totalClusters, message }')

# 5.5
doc.add_heading('5.5. submitGKPIdeas(formData)', level=3)
add_bullet('Викликається з: gkp_ideas.html')
add_bullet('Webhook: POST /webhook/gkp-ideas')
add_bullet('Payload: { doc_name, language, geo_target, source_spreadsheet_id }')
add_bullet('Відповідь: { success, spreadsheetUrl, totalKeywords, totalBatches, message }')

# 5.6
doc.add_heading('5.6. submitGKPMetrics(formData)', level=3)
add_bullet('Викликається з: gkp_metrics.html')
add_bullet('Webhook: POST /webhook/gkp-metrics')
add_bullet('Payload: { doc_name, language, geo_target, source_spreadsheet_id }')
add_bullet('Відповідь: { success, spreadsheetUrl, totalKeywords, keywordsWithData, keywordsNoData, message }')

# 5.7
doc.add_heading('5.7. submitPageSpeedTest(formData)', level=3)
add_bullet('Викликається з: pagespeed_form.html')
add_bullet('Webhook: POST /webhook/pagespeed-test')
add_bullet('Payload: { spreadsheetId, testMobile, testDesktop, batchSize: 5, delaySeconds: 5 }')
add_bullet('Відповідь: { success, spreadsheetUrl, totalUrls, processingTime, message }')

# 5.8
doc.add_heading('5.8. submitPdfAuditParse(formData)', level=3)
add_bullet('Викликається з: pdf_audit_form.html (картка "AI аналіз PDF документа")')
add_bullet('Webhook: POST /webhook-test/parse-pdf-audit')
add_bullet('Payload: { pdfUrl }')
add_bullet('Відповідь: { success, spreadsheetUrl, totalSheets, processingTime, message }')

# ═══════════════════════════════════════
# 6. N8N WORKFLOWS
# ═══════════════════════════════════════

doc.add_heading('6. n8n Воркфлоу — детальний опис', level=1)

# 6.1
doc.add_heading('6.1. SEO-аудит домену (основний)', level=2)
add_paragraph('Файл: Sheet 8 - Traffic Pages Nodes (1).json', italic=True, color=GRAY)

add_paragraph('Вебхуки:', bold=True)
add_table(
    ['Шлях', 'Метод', 'Призначення'],
    [
        ['seo-organic-traffic-v74', 'POST', 'Основна точка входу — запуск аудиту'],
        ['browse-ai-callback-v74', 'POST', 'Зворотний виклик від Browse AI після завершення скрапінгу'],
        ['seo-audit-full-v74', 'POST', 'Внутрішній тригер для збору даних Ahrefs + Serpstat'],
    ]
)

add_paragraph('Потік даних:', bold=True)
add_bullet('1. Отримує домен -> створює Google Sheet -> переміщує у папку на Drive')
add_bullet('2. Запускає Browse AI для скрапінгу SimilarWeb (трафік, канали, пристрої, соцмережі)')
add_bullet('3. Відразу повертає URL таблиці (асинхронна обробка)')
add_bullet('4. Browse AI callback записує поведінкові дані у Лист 1')
add_bullet('5. Запускає Full Audit внутрішнім вебхуком')
add_bullet('6. Паралельні виклики Ahrefs API (через MCP): трафік, беклінки, DR, реф-домени, анкори')
add_bullet('7. Виклики Serpstat API: ключові слова з позиціями (до 5000, з пагінацією)')

add_paragraph('Згенеровані листи:', bold=True)
add_table(
    ['Лист №', 'Назва', 'Джерело даних', 'Зміст'],
    [
        ['1', 'Органічний_трафік', 'Ahrefs + SimilarWeb', 'Огляд трафіку, історія, гео-розподіл, канали'],
        ['2', '(вбудований у Лист 1)', 'SimilarWeb / Browse AI', 'Брендовий vs небрендовий трафік, метрики залученості'],
        ['3', 'Посилальний_профіль', 'Ahrefs', 'Історія DR, реф-домени, розподіл анкорів'],
        ['4', 'Всі_беклінки', 'Ahrefs', 'Усі беклінки (топ-100 за DR)'],
        ['5', 'Топ_сторінки_за_посиланнями', 'Ahrefs', 'Топ сторінок за кількістю реф-доменів'],
        ['6', 'Поведінкові_метрики', 'SimilarWeb / Browse AI', 'Соцтрафік, реферери, видавці'],
        ['7', 'Ключові_фрази', 'Serpstat', 'До 5000 ключів з обсягами, позиціями, інтентом'],
        ['8', 'Трафікогенеруючі_сторінки', 'Serpstat (агреговано)', 'Сторінки згруповані за трафіком з топ-ключем'],
    ],
    col_widths=[1.5, 4.5, 3.5, 7]
)

add_paragraph('Зовнішні API:', bold=True)
add_bullet('Ahrefs (MCP протокол) — метрики домену, беклінки, анкори, DR')
add_bullet('SimilarWeb (через Browse AI скрапінг) — поведінкові метрики')
add_bullet('Serpstat API (api.serpstat.com/v4) — ключові слова')
add_bullet('Cloudinary — хостинг SVG-графіків')
add_bullet('Google Drive / Sheets API — управління файлами')

add_paragraph('Вхідні параметри:', bold=True)
add_code_block('{ "domain": "example.com (обов\'язково)", "country": "ua (за замовч.)", "top_pages_limit": 50 }')

# 6.2
doc.add_heading('6.2. Мастер-оркестратор "Аналіз домену"', level=2)
add_paragraph('Файл: Analiz_Domenu_Master.json', italic=True, color=GRAY)
add_paragraph('Webhook: POST /webhook/analiz-domenu', italic=True, color=GRAY)

add_paragraph('Потік даних:', bold=True)
add_bullet('1. Отримує email менеджера + опціональні блоки аналізу')
add_bullet('2. Шукає / створює персональну папку менеджера на Google Drive')
add_bullet('3. Створює підпапку проекту "{домен} - {дата}"')
add_bullet('4. Послідовно запускає обрані блоки:')
add_bullet('Аудит домену клієнта (викликає seo-organic-traffic-v74)', level=1)
add_bullet('Аудит конкурентів (цикл з затримкою 5 сек на кожного)', level=1)
add_bullet('Розширення семантики (викликає gkp-ideas)', level=1)
add_bullet('Збір метрик (викликає gkp-metrics)', level=1)
add_bullet('PageSpeed тест (викликає pagespeed-test)', level=1)
add_bullet('5. Переміщує всі згенеровані файли у папку проекту')
add_bullet('6. Повертає URL папки + статус кожного блоку')

# 6.3
doc.add_heading('6.3. AI-генератор звіту', level=2)
add_paragraph('Файл: SEO_Audit_AI_Report.json', italic=True, color=GRAY)
add_paragraph('Webhook: POST /webhook-test/seo-audit-ai-report', italic=True, color=GRAY)

add_paragraph('Потік даних:', bold=True)
add_bullet('1. Отримує URL / ID таблиці')
add_bullet('2. Зчитує всі 6 листів даних через Google Sheets API batchGet')
add_bullet('3. Запускає 6 паралельних GPT-4o агентів — кожен аналізує один лист:')
add_bullet('Агент 1: Аналіз органічного трафіку', level=1)
add_bullet('Агент 2: Аналіз посилального профілю', level=1)
add_bullet('Агент 3: Топ сторінки за посиланнями', level=1)
add_bullet('Агент 4: Поведінкові метрики', level=1)
add_bullet('Агент 5: Аналіз ключових слів', level=1)
add_bullet('Агент 6: Трафікогенеруючі сторінки', level=1)
add_bullet('4. Об\'єднує результати всіх 6 агентів')
add_bullet('5. Запускає Summary Agent — створює резюме з усіх секцій')
add_bullet('6. Створює Google Doc з форматуванням (шрифти Montserrat/Open Sans, navy/gold кольори)')
add_bullet('7. Переміщує документ до папки таблиці-джерела')
add_paragraph('Результат: URL Google Doc з форматованим AI-звітом', bold=True)

# 6.4
doc.add_heading('6.4. PDF Audit Parser', level=2)
add_paragraph('Файл: PDF_Audit_Parser.json', italic=True, color=GRAY)
add_paragraph('Webhook: POST /webhook-test/parse-pdf-audit', italic=True, color=GRAY)

add_paragraph('Потік даних:', bold=True)
add_bullet('1. Отримує URL PDF-файлу з Google Drive')
add_bullet('2. Завантажує PDF з Google Drive')
add_bullet('3. Завантажує PDF до OpenAI Files API')
add_bullet('4. Надсилає до GPT-4o зі структурованим промптом (temperature: 0.1)')
add_bullet('5. Видобуває 15 секцій: General Info, Summary, Content Types, URL Structure, Server Codes, Indexability, Meta Robots, Canonical, URL Depth, Loading Speed, Protocols, SEO Elements, Content Metrics, Errors, Scan Settings')
add_bullet('6. Створює нову Google-таблицю з відформатованими даними')
add_bullet('7. Застосовує форматування: секційні заголовки (темно-сині), підзаголовки (блакитні), заголовки таблиць (сірі)')
add_bullet('8. Видаляє тимчасовий файл з OpenAI')
add_bullet('9. Повертає URL таблиці')

add_paragraph('Вхід: { "pdfUrl": "Google Drive URL" }', bold=True)
add_paragraph('Вихід: { "spreadsheetUrl", "totalRows", "processingTime" }', bold=True)

# 6.5
doc.add_heading('6.5. GKP Universal System (семантика)', level=2)
add_paragraph('Файл: GKP_Universal_System.json', italic=True, color=GRAY)

add_table(
    ['Webhook', 'Призначення'],
    [
        ['gkp-ideas', 'Етап 1 — генерація ідей ключових слів із seed-фраз'],
        ['gkp-metrics', 'Етап 2 — отримання обсягів пошуку для конкретних ключів'],
    ]
)

add_paragraph('Етап 1 (Генерація ідей):', bold=True)
add_bullet('Зчитує seed-фрази з Google Sheet (колонка A) або тіла запиту')
add_bullet('Розбиває на батчі по 10 фраз')
add_bullet('Викликає Google Ads API generateKeywordIdeas з пагінацією (1000/сторінка)')
add_bullet('Дедуплікує, сортує за обсягом пошуку')
add_bullet('Записує у нову таблицю: [Ключове слово, Обсяг пошуку, Конкуренція, Індекс конкуренції, Seed-фраза]')

add_paragraph('Етап 2 (Метрики):', bold=True)
add_bullet('Зчитує ключові слова з Google Sheet або тіла запиту (до 100 000)')
add_bullet('Розбиває на батчі по 10 000')
add_bullet('Викликає Google Ads API generateKeywordHistoricalMetrics')
add_bullet('Записує: [Ключове слово, Обсяг, Конкуренція, Індекс, Помісячні обсяги]')

# 6.6
doc.add_heading('6.6. PageSpeed Test', level=2)
add_paragraph('Файл: PageSpeed_Test.json', italic=True, color=GRAY)
add_paragraph('Webhook: POST /webhook/pagespeed-test', italic=True, color=GRAY)

add_paragraph('Потік даних:', bold=True)
add_bullet('1. Зчитує URL-адреси з Google Sheet (колонка A)')
add_bullet('2. Створює нову таблицю з 4 вкладками: Mobile Results, Desktop Results, Comparison, Recommendations')
add_bullet('3. Тестує кожен URL через Google PageSpeed Insights API v5 (батчі по 5, затримка 5 сек)')
add_bullet('4. Метрики: Performance, Accessibility, Best Practices, SEO, FCP, LCP, TBT, CLS, Speed Index, TTI')
add_bullet('5. Запускає GPT-4o-mini для AI-коментарів до кожної вкладки')
add_bullet('6. Записує результати з форматуванням')

# 6.7
doc.add_heading('6.7. Аудит посилального профілю', level=2)
add_paragraph('Файл: Zovnishnya_skladova.json', italic=True, color=GRAY)
add_paragraph('Webhook: POST /webhook/link-profile-audit', italic=True, color=GRAY)

add_paragraph('Потік даних:', bold=True)
add_bullet('1. Створює папку на Google Drive')
add_bullet('2. Створює Google Sheet з 4 вкладками')
add_bullet('3. Паралельні виклики Ahrefs API (MCP): DR, статистика беклінків, анкори (топ-50), топ сторінки (100), всі беклінки (до 5000)')
add_bullet('4. Записує та форматує 4 листи:')
add_bullet('Link Profile (DR, ранг, статистика беклінків)', level=1)
add_bullet('Anchor List (топ-50 анкорів з %)', level=1)
add_bullet('Top Pages (100 сторінок за реф-доменами)', level=1)
add_bullet('All Backlinks (до 5000, відсортовано за трафіком)', level=1)

# ═══════════════════════════════════════
# 7. ENDPOINTS TABLE
# ═══════════════════════════════════════

doc.add_heading('7. Зведена таблиця ендпоінтів', level=1)

add_paragraph(
    'Базовий URL: https://n8n.rnd.webpromo.tools', bold=True
)

add_table(
    ['Ендпоінт', 'Метод', 'Воркфлоу', 'GAS-функція'],
    [
        ['/webhook/seo-organic-traffic-v74', 'POST', 'Sheet 8 - Traffic Pages', 'submitAudit()'],
        ['/webhook/browse-ai-callback-v74', 'POST', 'Sheet 8 - Traffic Pages', '(внутрішній callback)'],
        ['/webhook/seo-audit-full-v74', 'POST', 'Sheet 8 - Traffic Pages', '(внутрішній тригер)'],
        ['/webhook/analiz-domenu', 'POST', 'Analiz_Domenu_Master', 'submitAnalizDomenu()'],
        ['/webhook-test/seo-audit-ai-report', 'POST', 'SEO_Audit_AI_Report', 'submitAIAnalysis()'],
        ['/webhook-test/parse-pdf-audit', 'POST', 'PDF_Audit_Parser', 'submitPdfAuditParse()'],
        ['/webhook/gkp-ideas', 'POST', 'GKP_Universal_System', 'submitGKP(), submitGKPIdeas()'],
        ['/webhook/gkp-metrics', 'POST', 'GKP_Universal_System', 'submitGKPMetrics()'],
        ['/webhook/pagespeed-test', 'POST', 'PageSpeed_Test', 'submitPageSpeedTest()'],
        ['/webhook/link-profile-audit', 'POST', 'Zovnishnya_skladova', '(немає GAS-форми)'],
    ],
    col_widths=[5.5, 1.5, 4.5, 5]
)

add_paragraph(
    'Увага: /webhook-test/ ендпоінти — це тестові вебхуки n8n (активні лише коли воркфлоу '
    'відкрито в редакторі). /webhook/ — продакшн вебхуки (активні коли воркфлоу активовано).',
    bold=True, color=RGBColor(0xCC, 0x00, 0x00)
)

# ═══════════════════════════════════════
# 8. API CREDENTIALS
# ═══════════════════════════════════════

doc.add_heading('8. Зовнішні API та авторизація', level=1)

add_table(
    ['Сервіс', 'Тип авторизації', 'Деталі'],
    [
        ['Ahrefs', 'Bearer Token (MCP)', 'Токен у сховищі n8n credentials'],
        ['Google Ads', 'OAuth2 + Developer Token', 'Customer: 3460470607, Login: 3993420980'],
        ['Serpstat', 'HTTP Query Auth', 'Через сховище n8n credentials'],
        ['OpenAI', 'API Key', 'Моделі: gpt-4o (звіти), gpt-4o-mini (PageSpeed)'],
        ['Google Drive/Sheets/Docs', 'OAuth2', 'Через сховище n8n credentials'],
        ['PageSpeed Insights', 'API Key', 'Ключ у налаштуваннях воркфлоу'],
        ['Browse AI', 'HTTP Header Auth', 'Через сховище n8n credentials'],
        ['Cloudinary', 'API credentials', 'Для завантаження SVG-графіків'],
    ],
    col_widths=[4, 4, 8.5]
)

# ═══════════════════════════════════════
# 9. DRIVE STRUCTURE
# ═══════════════════════════════════════

doc.add_heading('9. Структура Google Drive', level=1)

drive_lines = [
    'Root: SEO - аудит (1A3Ak929G1c4XmZpPtI2FP4glrFE2-Bx2)',
    '|',
    '+-- [manager@email.com]/',
    '|   +-- [domain - 2026-02-16]/',
    '|       +-- SEO Audit - domain - date.xlsx    (8 листів)',
    '|       +-- GKP Ideas - date.xlsx',
    '|       +-- GKP Metrics - date.xlsx',
    '|       +-- PageSpeed Report - date.xlsx',
    '|       +-- SEO Audit AI Report - domain.gdoc',
    '|',
    '+-- Link Profile (окрема папка: 1VAwG8CkWIkhY...)',
    '    +-- [domain - date]/',
    '        +-- Link Profile Report.xlsx            (4 листи)',
]
for line in drive_lines:
    add_code_block(line)

# ═══════════════════════════════════════
# 10. FILE STRUCTURE
# ═══════════════════════════════════════

doc.add_heading('10. Структура файлів у репозиторії', level=1)

files_lines = [
    'n8n_seo_audit/',
    '|-- .mcp.json                              # MCP-сервер конфіг для n8n',
    '|-- gas/',
    '|   |-- Код.gs                             # Бекенд: маршрутизація + webhook-проксі',
    '|   |-- form.html                          # Сторінка: "Аналіз домену"',
    '|   |-- analiz_domenu_form.html            # Сторінка: "Мастер" оркестратор',
    '|   |-- gkp_form.html                      # Сторінка: GKP legacy',
    '|   |-- gkp_ideas.html                     # Сторінка: GKP Етап 1',
    '|   |-- gkp_metrics.html                   # Сторінка: GKP Етап 2',
    '|   |-- pagespeed_form.html                # Сторінка: PageSpeed тест',
    '|   |-- pdf_audit_form.html                # Сторінка: AI-аналіз + PDF парсер',
    '|',
    '|-- Analiz_Domenu_Master.json              # n8n: Мастер-оркестратор',
    '|-- SEO_Audit_AI_Report.json               # n8n: AI-генератор звіту (GPT-4o)',
    '|-- PDF_Audit_Parser.json                  # n8n: PDF -> таблиця парсер',
    '|-- GKP_Universal_System.json              # n8n: Google Keyword Planner',
    '|-- PageSpeed_Test.json                    # n8n: PageSpeed Insights тест',
    '|-- Sheet 8 - Traffic Pages Nodes (1).json # n8n: Основний SEO-аудит',
    '|-- Zovnishnya_skladova.json               # n8n: Аудит посилального профілю',
    '|-- To GKP.json                            # n8n: Тестовий воркфлоу',
]
for line in files_lines:
    add_code_block(line)

# ═══════════════════════════════════════
# 11. DEPLOYMENT NOTES
# ═══════════════════════════════════════

doc.add_heading('11. Нотатки щодо деплою', level=1)

doc.add_heading('GAS деплой', level=3)
add_bullet('Код знаходиться в редакторі GAS (посилання у розділі 2)')
add_bullet('Git-репозиторій є дзеркалом GAS-файлів для контролю версій')
add_bullet('Для деплою: скопіювати файли з gas/ у GAS-редактор -> Deploy as Web App')
add_bullet('URL Web App вбудований як iframe у Google Sites сторінку')

doc.add_heading('n8n деплой', level=3)
add_bullet('JSON-файли в репозиторії — це експорти n8n воркфлоу')
add_bullet('Для деплою: імпортувати JSON у n8n інстанс (n8n.rnd.webpromo.tools)')
add_bullet('Після імпорту: налаштувати credentials, активувати воркфлоу')

add_paragraph(
    'ВАЖЛИВО: /webhook-test/ URL працюють ЛИШЕ коли воркфлоу відкрито в редакторі n8n. '
    'Для продакшну потрібно змінити на /webhook/ та активувати воркфлоу.',
    bold=True, color=RGBColor(0xCC, 0x00, 0x00)
)

doc.add_heading('Поточний стан webhook-test', level=3)
add_paragraph(
    'Дві функції (submitAIAnalysis, submitPdfAuditParse) зараз вказують на webhook-test URL. '
    'Це означає, що вони працюватимуть тільки коли відповідні воркфлоу відкриті в редакторі n8n. '
    'Для продакшну: змінити на /webhook/ та активувати воркфлоу.'
)

# ── SAVE ──
output_path = '/home/user/n8n_seo_audit/SEO_Audit_Platform_Documentation.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
