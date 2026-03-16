"""Generate DOCX instruction document for Comparing SimilarWeb."""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# -- Style setup --
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def bold_run(paragraph, text):
    r = paragraph.add_run(text)
    r.bold = True
    return r


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        bold_run(p, bold_prefix)
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    if bold_prefix:
        bold_run(p, bold_prefix)
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_faq(question, answer_lines):
    p = doc.add_paragraph()
    bold_run(p, f'Q: {question}')
    if isinstance(answer_lines, str):
        answer_lines = [answer_lines]
    for line in answer_lines:
        ap = doc.add_paragraph()
        ap.add_run(f'A: {line}' if line == answer_lines[0] else line)
        ap.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Comparing SimilarWeb — Інструкція', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    run.font.size = Pt(26)

doc.add_paragraph()

# ============================================================
# 1. GENERAL DESCRIPTION
# ============================================================
doc.add_heading('1. Загальний опис', level=1)

doc.add_paragraph(
    'Система автоматичного збору та порівняння даних SimilarWeb '
    'для клієнтських сайтів та їхніх конкурентів.'
)

p = doc.add_paragraph()
bold_run(p, 'Стек: ')
p.add_run('Google Apps Script (панель управління) → Supabase (база даних) → '
          'n8n (автоматизація) → Browse AI (скрапінг SimilarWeb)')

doc.add_heading('Що робить система:', level=3)
add_bullet('Збирає дані SimilarWeb за останні 13 місяців для клієнтів та їхніх конкурентів')
add_bullet('Автоматично формує чергу задач, розбиваючи сайти на чанки по 5')
add_bullet('Запускає 3 Browse AI роботи паралельно для кожного чанку (Performance, Marketing Channels, AI Traffic)')
add_bullet('Зберігає дані в Supabase з можливістю порівняння в таблиці')
add_bullet('Сповіщає в Telegram про статус збору та помилки')

doc.add_heading('Метрики, які збираються:', level=3)
add_table(
    ['Група', 'Метрики'],
    [
        ['Performance', 'Monthly Visits, Unique Visitors, Visit Duration, Pages/Visit, Bounce Rate, Visits/Visitor, Deduplicated Audience, Page Views'],
        ['Marketing Channels', 'Direct, Organic Search, Paid Search, Display Ads, Social, Email'],
        ['AI Traffic', 'AI Traffic (%)'],
    ]
)

doc.add_heading('Воркфлоу в n8n:', level=3)
add_table(
    ['Воркфлоу', 'Що робить', 'Розклад'],
    [
        ['Comparing SimilarWeb — supa base', 'Обробка черги задач + збір даних', 'Кожні 7 хвилин'],
        ['Comparing SimilarWeb — supa base\n(Проверка полноты)', 'Аналіз пропусків і створення нових задач', '2 рази на день\n(08:00, 20:00 UTC)'],
        ['Retry Stale Tasks', 'Перезапуск зависших задач (>10 хв в processing)', 'Кожні 15 хвилин'],
    ]
)

# ============================================================
# 2. STEP-BY-STEP GUIDE
# ============================================================
doc.add_heading('2. Інструкція до застосування', level=1)

# Step 1
doc.add_heading('Крок 1. Налаштування підключення до Supabase', level=2)
add_numbered('Відкрийте панель управління (GAS web app)')
add_numbered('Перейдіть на вкладку Settings')
p = doc.add_paragraph(style='List Number')
p.add_run('Введіть:')
add_bullet('Supabase URL', bold_prefix=None)
add_bullet('Supabase Key — Anon Key (запитати у адміністратора)', bold_prefix=None)
add_numbered('Натисніть Test Connection')
add_numbered('Якщо з\'явилося повідомлення з кількістю знайдених клієнтів — підключення успішне')

# Step 2
doc.add_heading('Крок 2. Додавання клієнта', level=2)
add_numbered('Перейдіть на вкладку Clients')
add_numbered('Введіть домен клієнта (наприклад: example.com) та email відповідального')
add_numbered('Натисніть Add Client')
add_numbered('Домен автоматично нормалізується (видаляються http://, www., тощо)')

# Step 3
doc.add_heading('Крок 3. Додавання конкурентів', level=2)
add_numbered('На вкладці Clients знайдіть потрібного клієнта')
add_numbered('Натисніть кнопку Competitors')
add_numbered('У модальному вікні введіть домен конкурента і натисніть Add')
add_numbered('Повторіть для всіх конкурентів')
add_numbered('Максимум — без обмежень, але система розбиває на чанки по 5 сайтів')

# Step 4
doc.add_heading('Крок 4. Запуск збору даних', level=2)

p = doc.add_paragraph()
bold_run(p, 'Автоматично: ')
p.add_run('система двічі на день (08:00 та 20:00 UTC) перевіряє повноту даних '
          'і створює задачі для відсутніх періодів.')

p = doc.add_paragraph()
bold_run(p, 'Вручну:')
add_numbered('Перейдіть на вкладку Dashboard')
add_numbered('Натисніть Build Full Queue')
add_numbered('Система проаналізує пропуски за останні 13 місяців і створить задачі')

doc.add_paragraph()
p = doc.add_paragraph()
bold_run(p, 'Після створення задач:')
add_bullet('Кожні 7 хвилин n8n бере з черги до 9 задач')
add_bullet('Між задачами — рандомна затримка 3–7 хвилин (захист від блокування Browse AI)')
add_bullet('Для кожної задачі паралельно запускаються 3 роботи')
add_bullet('Результати парсяться та зберігаються в Supabase')

# Step 5
doc.add_heading('Крок 5. Моніторинг процесу', level=2)

p = doc.add_paragraph()
bold_run(p, 'Dashboard:')
add_bullet('Картки зі статусами: pending (очікує), processing (обробляється), done (готово), error (помилка)')

doc.add_paragraph()
p = doc.add_paragraph()
bold_run(p, 'Queue (вкладка):')
add_bullet('Таблиця останніх 100 задач з деталями: клієнт, період, сайти, статус, дата створення')

doc.add_paragraph()
p = doc.add_paragraph()
bold_run(p, 'Logs (вкладка):')
add_bullet('Логи помилок з фільтрацією (всі / невирішені / вирішені)')
add_bullet('Кнопка Retry для повторного запуску проблемної задачі')
add_bullet('Кнопка Resolve для позначення помилки як вирішеної')

doc.add_paragraph()
p = doc.add_paragraph()
bold_run(p, 'Telegram:')
add_bullet('Автоматичні сповіщення про завершення збору для клієнта, помилки роботів, зависші задачі')

# Step 6
doc.add_heading('Крок 6. Перегляд результатів', level=2)
add_numbered('Перейдіть на вкладку Reviews')
add_numbered('Оберіть клієнта зі списку')
add_numbered('Оберіть діапазон місяців')
p = doc.add_paragraph(style='List Number')
p.add_run('У таблиці порівняння:')
add_bullet('Рядки клієнта виділені синім')
add_bullet('Рядки конкурентів — сірим')
add_bullet('Всі метрики: трафік, поведінкові, канали, AI Traffic')

# ============================================================
# 3. FAQ
# ============================================================
doc.add_heading('3. FAQ', level=1)

# -- General --
doc.add_heading('Загальні питання', level=2)

add_faq(
    'Скільки часу займає повний збір даних для нового клієнта?',
    [
        'Залежить від кількості конкурентів.',
        'Формула: ceil((1 клієнт + N конкурентів) / 5) × 13 місяців × ~5-10 хв на задачу.',
        'Для клієнта з 4 конкурентами (1 чанк × 13 місяців = 13 задач) — приблизно 1–2 години.',
    ]
)

add_faq(
    'Чому дані збираються за 13 місяців?',
    ['Щоб забезпечити порівняння "місяць до місяця" (YoY) для останнього місяця.']
)

add_faq(
    'Що таке "чанк"?',
    [
        'Група до 5 сайтів, які подаються в один запит Browse AI.',
        'Наприклад, якщо у клієнта 8 конкурентів, буде 2 чанки: '
        '[клієнт + 4 конкуренти] та [клієнт + 4 конкуренти].',
    ]
)

# -- Errors --
doc.add_heading('Помилки та відновлення', level=2)

add_faq(
    'Задача зависла у статусі "processing" — що робити?',
    [
        'Нічого. Воркфлоу "Retry Stale Tasks" автоматично скидає задачі, '
        'що обробляються більше 10 хвилин. Задача повернеться в статус '
        '"pending" і буде оброблена повторно.',
    ]
)

add_faq(
    'У логах з\'явилася помилка робота — що це означає?',
    [
        'Browse AI не зміг зібрати дані зі сторінки SimilarWeb. Можливі причини:',
        '• SimilarWeb змінив інтерфейс',
        '• Сайт не має даних у SimilarWeb',
        '• Тимчасова помилка завантаження',
        '• Вичерпано ліміт Browse AI',
        '',
        'Дія: натисніть Retry в логах. Якщо помилка повторюється — '
        'перевірте, чи є дані на SimilarWeb вручну.',
    ]
)

add_faq(
    'Задачі створюються, але не обробляються — чому?',
    [
        'Перевірте:',
        '1. Воркфлоу "Comparing SimilarWeb — supa base" активний у n8n',
        '2. Browse AI credential дійсний',
        '3. Ліміт Browse AI не вичерпаний (макс. 9 задач за цикл)',
    ]
)

add_faq(
    'Як перезапустити збір для конкретного місяця?',
    [
        'Знайдіть задачу у вкладці Queue або Logs і натисніть Retry.',
        'Або видаліть задачу і запустіть Build Full Queue — система сама знайде пропуск.',
    ]
)

# -- Clients --
doc.add_heading('Управління клієнтами', level=2)

add_faq(
    'Можна видалити клієнта?',
    [
        'Так, через вкладку Clients. Видалення — м\'яке (soft delete): '
        'клієнт і його конкуренти позначаються як видалені, але дані зберігаються в базі.',
    ]
)

add_faq(
    'Можна додати конкурента після того, як збір вже запущений?',
    [
        'Так. Додайте конкурента, потім натисніть Build Full Queue. '
        'Система створить задачі тільки для нового конкурента '
        '(існуючі дані не перезаписуються).',
    ]
)

add_faq(
    'Що станеться, якщо додати домен з www або http://?',
    [
        'Система автоматично нормалізує домен — видалить протокол, www, '
        'слеш в кінці та приведе до нижнього регістру.',
        'Приклад: https://www.Example.com/ → example.com',
    ]
)

# -- Data --
doc.add_heading('Дані та звіти', level=2)

add_faq(
    'Чому для деяких сайтів немає даних?',
    [
        'SimilarWeb може не мати даних для малих сайтів або нових доменів. '
        'Перевірте наявність даних на pro.similarweb.com вручну.',
    ]
)

add_faq(
    'Чи оновлюються дані автоматично щомісяця?',
    ['Так. Двічі на день система перевіряє повноту і додає задачі для нових періодів.']
)

add_faq(
    'Як зрозуміти, що збір для клієнта завершений?',
    [
        'На вкладці Clients біля клієнта з\'являється статус "done" '
        'та мітка часу останнього завершення. Також приходить сповіщення в Telegram.',
    ]
)

# -- Save --
output_path = '/home/user/n8n_seo_audit/Comparing SimilarWeb/Comparing_SimilarWeb_Instruction.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
