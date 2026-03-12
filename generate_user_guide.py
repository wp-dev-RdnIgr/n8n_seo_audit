#!/usr/bin/env python3
"""Generate SimilarWeb Comparing user guide in DOCX format (Ukrainian)."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# -- Styles --
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.paragraph_format.space_after = Pt(6)
style_normal.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)


def add_tip(text, icon="💡"):
    p = doc.add_paragraph()
    run = p.add_run(f"{icon} {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x06, 0x5F, 0x46)
    run.font.italic = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_warning(text):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️ {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    run.font.italic = True


def add_step(number, text):
    p = doc.add_paragraph()
    run_num = p.add_run(f"Крок {number}. ")
    run_num.bold = True
    run_num.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    p.add_run(text)


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f" — {text}")
    else:
        p.add_run(text)


# ============================================================
# TITLE PAGE
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SimilarWeb Comparing")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
run.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Інструкція користувача")
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(20)
run3 = p3.add_run("Система аналізу та порівняння трафіку конкурентів\nна базі даних SimilarWeb")
run3.font.size = Pt(12)
run3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.paragraph_format.space_before = Pt(40)
run4 = p4.add_run("Версія 1.0 | Березень 2026")
run4.font.size = Pt(10)
run4.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading("Зміст", level=1)
toc_items = [
    "1. Загальний огляд сервісу",
    "2. Навігація по інтерфейсу",
    "3. Початок роботи: додавання клієнта",
    "4. Управління конкурентами",
    "5. Збір даних: черга задач",
    "6. Відстеження статусу оновлення",
    "7. Dashboard — головна сторінка",
    "8. Аналіз клієнта (Reviews)",
    "   8.1. Client Overview — огляд метрик",
    "   8.2. Competitor Comparison — порівняння конкурентів",
    "   8.3. Traffic Sources — джерела трафіку",
    "9. Черга задач (Queue)",
    "10. Журнал помилок (Logs)",
    "11. Налаштування (Settings)",
    "12. FAQ — часті питання",
    "13. Що можна і чого не можна робити",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith("   "):
        p.paragraph_format.left_indent = Cm(1)

doc.add_page_break()

# ============================================================
# 1. OVERVIEW
# ============================================================
doc.add_heading("1. Загальний огляд сервісу", level=1)

doc.add_paragraph(
    "SimilarWeb Comparing — це внутрішній інструмент для аналізу та порівняння "
    "трафіку клієнтських сайтів із конкурентами. Система автоматично збирає дані "
    "з SimilarWeb, зберігає їх у базу даних Supabase та надає зручний інтерфейс "
    "для аналізу.")

doc.add_heading("Основні можливості:", level=3)
add_bullet("Додавання клієнтів та їх конкурентів")
add_bullet("Автоматичний збір даних SimilarWeb через n8n-автоматизацію")
add_bullet("Порівняння трафіку за місяцями (Month-over-Month) та роками (Year-over-Year)")
add_bullet("Детальне порівняння конкурентів за 15 метриками")
add_bullet("Аналіз джерел трафіку з графіками та таблицями")
add_bullet("Відстеження статусу збору даних в реальному часі")
add_bullet("Журнал помилок з можливістю повторного запуску")

doc.add_heading("Доступні метрики (15):", level=3)
metrics = [
    ("Monthly Visits", "загальна кількість візитів за місяць"),
    ("Unique Visitors", "унікальні відвідувачі"),
    ("Visits per Visitor", "середня кількість візитів на відвідувача"),
    ("Deduplicated Audience", "дедуплікована аудиторія"),
    ("Page Views", "загальна кількість переглядів сторінок"),
    ("Visit Duration", "середня тривалість візиту"),
    ("Pages per Visit", "середня кількість сторінок за візит"),
    ("Bounce Rate", "показник відмов"),
    ("Direct", "прямий трафік"),
    ("Organic Search", "органічний пошуковий трафік"),
    ("Paid Search", "платний пошуковий трафік"),
    ("Display Ads", "медійна реклама"),
    ("Social", "трафік з соціальних мереж"),
    ("Email", "трафік з email-розсилок"),
    ("AI Traffic", "трафік від AI-сервісів"),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light List Accent 1'
table.columns[0].width = Cm(4)
table.columns[1].width = Cm(12)
hdr = table.rows[0].cells
hdr[0].text = "Метрика"
hdr[1].text = "Опис"
for m, d in metrics:
    row = table.add_row().cells
    row[0].text = m
    row[1].text = d

doc.add_page_break()

# ============================================================
# 2. NAVIGATION
# ============================================================
doc.add_heading("2. Навігація по інтерфейсу", level=1)

doc.add_paragraph(
    "Інтерфейс складається з верхньої панелі навігації з 6 вкладками:")

tabs = [
    ("Dashboard", "Головна сторінка зі статистикою та швидким доступом до клієнтів"),
    ("Clients", "Управління клієнтами: додавання, видалення, управління конкурентами"),
    ("Reviews", "Детальний аналіз кожного клієнта з графіками та порівняннями"),
    ("Queue", "Управління чергою задач на збір даних"),
    ("Logs", "Журнал помилок збору даних"),
    ("Settings", "Налаштування підключення до Supabase"),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Вкладка"
hdr[1].text = "Призначення"
for t, d in tabs:
    row = table.add_row().cells
    row[0].text = t
    row[1].text = d

doc.add_paragraph()
add_tip("На вкладках Queue та Clients відображаються лічильники — кількість задач/клієнтів.")
add_tip("У правому верхньому куті показано статус підключення до Supabase (зелена точка = підключено).")

doc.add_page_break()

# ============================================================
# 3. ADDING CLIENT
# ============================================================
doc.add_heading("3. Початок роботи: додавання клієнта", level=1)

doc.add_paragraph(
    "Перший крок для початку аналізу — додати клієнтський сайт у систему.")

doc.add_heading("Покрокова інструкція:", level=3)
add_step(1, 'Перейдіть на вкладку "Clients".')
add_step(2, 'У полі "Client Domain" введіть домен клієнта (наприклад: example.com.ua).')
add_step(3, 'У полі "Manager Contact" вкажіть email менеджера або Telegram-нік (наприклад: @username). Це поле необов\'язкове.')
add_step(4, 'Натисніть кнопку "+ Add".')

doc.add_paragraph()
add_tip("Домен автоматично нормалізується: видаляється https://, www., кінцевий слеш, приводиться до нижнього регістру.")
add_warning("Не додавайте один і той самий домен двічі — система відхилить дублікат.")

doc.add_heading("Що відбувається після додавання:", level=3)
add_bullet("Клієнт з'являється в таблиці з статусом «New»")
add_bullet("Тепер потрібно додати конкурентів (див. розділ 4)")
add_bullet("Після додавання конкурентів — запустити збір даних (див. розділ 5)")

doc.add_heading("Видалення клієнта:", level=3)
doc.add_paragraph(
    'Натисніть іконку кошика (🗑) у рядку клієнта. З\'явиться діалог підтвердження. '
    'Видалення є "м\'яким" (soft delete) — дані зберігаються в базі, але клієнт '
    'та всі його конкуренти зникають з інтерфейсу.')
add_warning("При видаленні клієнта всі його конкуренти також видаляються.")

doc.add_page_break()

# ============================================================
# 4. COMPETITORS
# ============================================================
doc.add_heading("4. Управління конкурентами", level=1)

doc.add_paragraph(
    "Кожному клієнту можна додати необмежену кількість конкурентів для порівняння.")

doc.add_heading("Додавання конкурентів:", level=3)
add_step(1, 'На вкладці "Clients" натисніть іконку 👥 (Manage Competitors) у рядку потрібного клієнта.')
add_step(2, "Відкриється модальне вікно зі списком поточних конкурентів.")
add_step(3, 'Введіть домен конкурента в поле (наприклад: competitor.com) та натисніть "+ Add".')
add_step(4, "Конкурент з'явиться у списку нижче.")

doc.add_paragraph()
add_tip("Рекомендована кількість конкурентів: 5–15 на клієнта. Занадто багато конкурентів сповільнюють збір даних.")
add_warning("Якщо конкурент раніше був видалений і ви додаєте його знову — система автоматично відновить його без помилок.")

doc.add_heading("Видалення конкурента:", level=3)
doc.add_paragraph(
    'У модальному вікні конкурентів натисніть іконку ✕ біля потрібного конкурента. '
    'Конкурент буде "м\'яко" видалений — його дані залишаться в базі, але він '
    'перестане відображатися в порівняннях.')

doc.add_page_break()

# ============================================================
# 5. DATA COLLECTION
# ============================================================
doc.add_heading("5. Збір даних: черга задач", level=1)

doc.add_paragraph(
    "Після додавання клієнта та конкурентів необхідно запустити збір даних SimilarWeb. "
    "Це відбувається через систему черги задач.")

doc.add_heading("Спосіб 1: Автоматичне створення задач (Build Queue)", level=3)
add_step(1, 'На Dashboard натисніть кнопку "Build Queue".')
add_step(2, "Система автоматично створить задачі для всіх клієнтів за відсутні періоди.")
add_step(3, "Задачі почнуть виконуватись автоматично через n8n-робота.")

doc.add_paragraph()
add_tip("Build Queue — найзручніший спосіб. Система сама визначить, які дані відсутні, та створить задачі тільки для них.")

doc.add_heading("Спосіб 2: Ручне створення задач", level=3)
add_step(1, 'Перейдіть на вкладку "Queue".')
add_step(2, "Оберіть клієнта зі списку.")
add_step(3, "Оберіть період (місяць).")
add_step(4, 'Натисніть "Create".')

doc.add_paragraph()
doc.add_paragraph(
    "Система розбиває сайти на групи по 5 (чанки) для паралельної обробки. "
    "Наприклад, клієнт із 12 конкурентами = 13 сайтів = 3 чанки.")

doc.add_heading("Статуси задач:", level=3)
statuses = [
    ("Pending", "Задача створена, очікує на обробку"),
    ("Processing", "Задача зараз виконується роботом"),
    ("Done", "Дані успішно зібрані"),
    ("Error", "Виникла помилка під час збору"),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Статус"
hdr[1].text = "Опис"
for s, d in statuses:
    row = table.add_row().cells
    row[0].text = s
    row[1].text = d

doc.add_page_break()

# ============================================================
# 6. STATUS TRACKING
# ============================================================
doc.add_heading("6. Відстеження статусу оновлення", level=1)

doc.add_paragraph(
    "На Dashboard та в таблицях клієнтів відображається статус збору даних:")

statuses_client = [
    ("🔄 Updating", "Синій бейдж з анімованою іконкою обертання", "Є задачі зі статусом Pending або Processing — дані збираються"),
    ("✅ Done", "Зелений бейдж з галочкою + дата/час", "Всі задачі завершені. Показується дата останнього завершення"),
    ("⏳ New", "Жовтий бейдж з годинником", "Клієнт щойно доданий, задачі ще не створювались"),
]
table = doc.add_table(rows=1, cols=3)
table.style = 'Light List Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Статус"
hdr[1].text = "Вигляд"
hdr[2].text = "Що означає"
for s, v, d in statuses_client:
    row = table.add_row().cells
    row[0].text = s
    row[1].text = v
    row[2].text = d

doc.add_paragraph()
add_tip("Статуси оновлюються при кожному натисканні Refresh. Вони обчислюються в реальному часі з черги задач.")

doc.add_page_break()

# ============================================================
# 7. DASHBOARD
# ============================================================
doc.add_heading("7. Dashboard — головна сторінка", level=1)

doc.add_paragraph(
    "Dashboard — це стартова сторінка із загальною статистикою та швидким доступом до основних функцій.")

doc.add_heading("Картки статистики:", level=3)
stats = [
    ("Clients", "Кількість активних клієнтів"),
    ("Competitors", "Загальна кількість конкурентів"),
    ("Pending", "Задачі, що очікують виконання"),
    ("Done", "Кількість завершених задач"),
    ("Errors", "Помилки (нерозв'язані / загальні)"),
    ("Data Rows", "Кількість записів даних SimilarWeb"),
]
for s, d in stats:
    add_bullet(d, s)

doc.add_heading("Кнопки швидких дій:", level=3)
add_bullet("Перехід на вкладку додавання клієнтів", "Add Client")
add_bullet("Перехід до черги задач", "Queue")
add_bullet("Оновлення статистики та таблиці клієнтів", "Refresh")
add_bullet("Автоматичне створення задач для збору відсутніх даних", "Build Queue")

doc.add_heading("Таблиця Client Reviews:", level=3)
doc.add_paragraph(
    "Показує всіх активних клієнтів зі статусом оновлення та кнопкою Review "
    "для переходу до детального аналізу.")

doc.add_page_break()

# ============================================================
# 8. REVIEWS
# ============================================================
doc.add_heading("8. Аналіз клієнта (Reviews)", level=1)

doc.add_paragraph(
    'Вкладка Reviews — основний інструмент аналізу. Оберіть клієнта зі списку або '
    'натисніть "Review" на Dashboard.')

doc.add_paragraph(
    "Сторінка аналізу має 3 під-вкладки:")

# 8.1
doc.add_heading("8.1. Client Overview — огляд метрик клієнта", level=2)

doc.add_heading("Month over Month (MoM):", level=3)
doc.add_paragraph(
    "Порівняння обраного місяця з попереднім. Показує зміну кожної метрики "
    "у відсотках із кольоровим індикатором (зелений = зростання, червоний = падіння).")
add_bullet("Оберіть місяць у випадаючому календарі")
add_bullet('Оберіть метрики (або залиште "All metrics")')
add_bullet('Натисніть "Apply"')

doc.add_heading("Year over Year (YoY):", level=3)
doc.add_paragraph(
    "Порівняння обраного місяця з аналогічним місяцем минулого року. "
    "Ідеально для оцінки сезонності та довгострокових трендів.")

doc.add_heading("All Metrics Data:", level=3)
doc.add_paragraph(
    "Повна таблиця даних за обраний діапазон дат. Можна обрати період "
    "«від-до» та фільтрувати за метриками.")

# 8.2
doc.add_heading("8.2. Competitor Comparison — порівняння конкурентів", level=2)

doc.add_heading("MoM Chart:", level=3)
doc.add_paragraph(
    "Стовпчиковий графік, що порівнює обрану метрику клієнта та всіх конкурентів "
    "за обраний місяць проти попереднього. Під графіком — детальна таблиця.")

doc.add_heading("YoY Chart:", level=3)
doc.add_paragraph(
    "Аналогічне порівняння, але за рік (обраний місяць проти того ж місяця минулого року).")

doc.add_heading("Summary Table:", level=3)
doc.add_paragraph(
    "Зведена таблиця усіх метрик: рядки — метрики, стовпці — сайти. "
    "Клієнтський сайт виділяється синім фоном. Можна обрати конкурентів "
    "для відображення через мультиселект.")
add_tip("Таблиця має горизонтальний скрол, якщо конкурентів багато.")

# 8.3
doc.add_heading("8.3. Traffic Sources — джерела трафіку", level=2)

doc.add_heading("Traffic & Engagement:", level=3)
doc.add_paragraph(
    "Горизонтальні бари трафіку кожного сайту та таблиця залученості "
    "(візити, унікальні відвідувачі, сторінки за візит, тривалість, bounce rate).")
add_tip("Переможець у кожній метриці позначається іконкою 🏆.")

doc.add_heading("Traffic Over Time:", level=3)
doc.add_paragraph(
    "Лінійний графік обраної метрики за весь доступний період. "
    "Показує тренди для клієнта та конкурентів одночасно.")

doc.add_heading("Channel Traffic:", level=3)
doc.add_paragraph(
    "7 вкладок каналів: Direct, Organic Search, Paid Search, Display Ads, Social, Email, AI Traffic. "
    "Кожна показує лінійний графік трафіку з цього каналу по всіх сайтах.")

doc.add_heading("AI Traffic Distribution:", level=3)
doc.add_paragraph(
    "Окремий графік порівняння AI-трафіку між усіма сайтами.")

doc.add_heading("Traffic Source Distribution:", level=3)
doc.add_paragraph(
    "Кругова діаграма (donut chart) розподілу джерел трафіку за обраний місяць. "
    "Показує відсотки кожного каналу.")

doc.add_heading("Source Trends Over Time:", level=3)
doc.add_paragraph(
    "Графік із заповненими областями — показує, як змінювалось співвідношення "
    "каналів трафіку з часом.")

doc.add_heading("Source Details:", level=3)
doc.add_paragraph(
    "Детальна таблиця числових значень трафіку з кожного каналу по кожному сайту. "
    "Переможці позначаються 🏆.")

doc.add_page_break()

# ============================================================
# 9. QUEUE
# ============================================================
doc.add_heading("9. Черга задач (Queue)", level=1)

doc.add_paragraph(
    "Вкладка Queue дозволяє керувати задачами збору даних.")

doc.add_heading("Створення задач:", level=3)
add_step(1, "Оберіть клієнта зі списку")
add_step(2, "Оберіть місяць для збору")
add_step(3, 'Натисніть "Create"')

doc.add_paragraph()
add_tip("Система автоматично пропускає вже існуючі задачі для цього клієнта/періоду.")

doc.add_heading("Фільтрація:", level=3)
doc.add_paragraph(
    "Використовуйте випадаючий список для фільтрації за статусом: All, Pending, Processing, Done, Error.")

doc.add_heading("Дії із задачами:", level=3)
add_bullet("Повторний запуск задачі зі статусом Error", "Retry")
add_bullet("Видалення задачі з черги (повне видалення)", "Delete")

doc.add_page_break()

# ============================================================
# 10. LOGS
# ============================================================
doc.add_heading("10. Журнал помилок (Logs)", level=1)

doc.add_paragraph(
    "Вкладка Logs відображає помилки, що виникли під час збору даних.")

doc.add_heading("Фільтри:", level=3)
add_bullet("Всі помилки (за замовчуванням)", "All errors")
add_bullet("Тільки нерозв'язані", "Unresolved")
add_bullet("Тільки розв'язані", "Resolved")

doc.add_heading("Інформація про помилку:", level=3)
add_bullet("Дата та час виникнення")
add_bullet("Тип робота (Performance / AI Traffic)")
add_bullet("Період збору")
add_bullet("Список сайтів, для яких виникла помилка")
add_bullet("Текст помилки")
add_bullet("Кількість повторних спроб")

doc.add_heading("Дії:", level=3)
add_bullet("Знаходить відповідну задачу та запускає її повторно. Лічильник повторів збільшується.", "Retry")
add_bullet("Позначає помилку як розв'язану з відміткою часу.", "Mark Resolved")

doc.add_page_break()

# ============================================================
# 11. SETTINGS
# ============================================================
doc.add_heading("11. Налаштування (Settings)", level=1)

doc.add_paragraph(
    "Вкладка Settings містить налаштування підключення до бази даних Supabase.")

doc.add_heading("Параметри:", level=3)
add_bullet("URL проєкту Supabase (формат: https://xxxxx.supabase.co)", "Supabase URL")
add_bullet("Ключ для авторизації API-запитів", "Service Role Key")

doc.add_heading("Дії:", level=3)
add_bullet("Зберігає налаштування", "Save")
add_bullet("Перевіряє підключення та доступ до таблиць", "Test Connection")

doc.add_paragraph()
add_warning("Не змінюйте ці налаштування без консультації з адміністратором!")

doc.add_page_break()

# ============================================================
# 12. FAQ
# ============================================================
doc.add_heading("12. FAQ — часті питання", level=1)

faqs = [
    (
        "Як довго збираються дані після створення задач?",
        "Зазвичай від 5 до 30 хвилин залежно від кількості конкурентів та навантаження. "
        "Слідкуйте за статусом на Dashboard — коли всі задачі виконані, з'явиться статус «Done»."
    ),
    (
        "Чому у деяких конкурентів відсутні дані?",
        "SimilarWeb може не мати даних для малих сайтів з невеликим трафіком. "
        "Також можливі тимчасові помилки збору — перевірте вкладку Logs."
    ),
    (
        "Що означає статус «New» у клієнта?",
        "Клієнт щойно доданий, але задачі на збір даних ще не створювались. "
        "Натисніть «Build Queue» на Dashboard або створіть задачі вручну на вкладці Queue."
    ),
    (
        "Можу я додати конкурента, якого раніше видалив?",
        "Так! Система автоматично відновить раніше видаленого конкурента замість створення дублікату."
    ),
    (
        "Чому дані не завантажуються на сторінці аналізу?",
        "Переконайтесь, що: 1) вибрано правильний місяць у фільтрі дат; "
        "2) для цього клієнта є зібрані дані за обраний період; "
        "3) статус підключення Supabase «зелений» у верхньому правому куті."
    ),
    (
        "Що станеться, якщо видалити клієнта?",
        "Клієнт та його конкуренти будуть приховані (soft delete). "
        "Зібрані дані SimilarWeb залишаться в базі. "
        "Адміністратор може відновити видаленого клієнта при потребі."
    ),
    (
        "Скільки конкурентів можна додати?",
        "Технічних обмежень немає, але рекомендується 5–15 конкурентів. "
        "Більша кількість збільшує час збору даних та ускладнює графіки."
    ),
    (
        "Як часто оновлюються дані?",
        "SimilarWeb оновлює дані щомісячно. Рекомендується запускати Build Queue "
        "на початку кожного місяця для збору даних за попередній місяць."
    ),
    (
        "Таблиця не поміщається на екрані — що робити?",
        "Таблиці з великою кількістю конкурентів мають горизонтальний скрол. "
        "Прокрутіть таблицю вправо, щоб побачити всіх конкурентів."
    ),
    (
        "Помилка при додаванні конкурента — «Competitor already exists»",
        "Цей конкурент вже додано до клієнта. Перевірте список конкурентів."
    ),
    (
        "Що означає іконка 🏆 у таблицях?",
        "Це позначення переможця — сайт із найкращим показником по даній метриці "
        "(найбільший трафік, найнижчий bounce rate тощо)."
    ),
]

for q, a in faqs:
    p_q = doc.add_paragraph()
    run_q = p_q.add_run(f"❓ {q}")
    run_q.bold = True
    run_q.font.size = Pt(11)

    p_a = doc.add_paragraph(a)
    p_a.paragraph_format.left_indent = Cm(0.5)
    p_a.paragraph_format.space_after = Pt(10)

doc.add_page_break()

# ============================================================
# 13. DO's and DON'Ts
# ============================================================
doc.add_heading("13. Що можна і чого не можна робити", level=1)

doc.add_heading("✅ Можна:", level=3)
dos = [
    "Додавати клієнтів та конкурентів у будь-який час",
    "Запускати Build Queue для автоматичного збору даних",
    "Повторно запускати задачі з помилками через Retry",
    "Видаляти клієнтів або конкурентів (дані зберігаються)",
    "Фільтрувати графіки за метриками та періодами",
    "Використовувати горизонтальний скрол у широких таблицях",
    "Порівнювати дані за будь-який доступний період",
    "Оновлювати статистику кнопкою Refresh",
]
for d in dos:
    add_bullet(d)

doc.add_heading("❌ Не можна:", level=3)
donts = [
    "Змінювати налаштування Supabase без узгодження з адміністратором",
    "Видаляти задачі зі статусом Processing (зачекайте завершення)",
    "Додавати невалідні домени (без крапки або з пробілами)",
    "Одночасно створювати дублікати задач для одного клієнта/періоду",
    "Очікувати моментальний збір даних — це може зайняти до 30 хвилин",
    "Закривати сторінку під час операцій (зачекайте завершення)",
]
for d in donts:
    add_bullet(d)

doc.add_heading("💡 Рекомендації:", level=3)
recs = [
    "Запускайте Build Queue раз на місяць для збору свіжих даних",
    "Перевіряйте вкладку Logs після кожного збору даних",
    "Використовуйте YoY порівняння для оцінки сезонності",
    "Додавайте прямих конкурентів клієнта для найкориснішого аналізу",
    "Зберігайте email менеджера при додаванні клієнта — це допоможе з ідентифікацією",
]
for r in recs:
    add_bullet(r)

# ============================================================
# FOOTER
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_footer.add_run("— WebPromo | SimilarWeb Comparing v1.0 —")
run_f.font.size = Pt(9)
run_f.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# Save
output_path = "/home/user/n8n_seo_audit/SimilarWeb_Comparing_User_Guide_UA.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
